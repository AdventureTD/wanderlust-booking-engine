"""Tests for the Word-template invoice renderer."""

from datetime import date
from pathlib import Path
import sys

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from booking_engine.invoice import Guest, Invoice
from booking_engine.invoice_word import (
    MAX_PAYMENTS,
    MAX_ROOMS,
    build_invoice_context,
    find_libreoffice,
    render_invoice_docx,
    render_invoice_word_pdf,
)


def make_invoice(room_count=2, payment_count=2):
    display_items = []
    for index in range(room_count):
        display_items.append({
            "label": f"Room {index + 1}",
            "tax_class": "accommodation",
            "quantity": 7,
            "unit_price": 1000,
            "net": 7000,
            "vat_rate": 0.10,
            "vat": 700,
            "gross": 7700,
            "room_quantity": 1,
        })
    payments = [
        {"datePaid": f"2026-08-{index + 1:02d}", "paymentAmount": 100 * (index + 1)}
        for index in range(payment_count)
    ]
    return Invoice.from_quote(
        invoice_number="WBE-TEST-001",
        issue_date=date(2026, 8, 17),
        guest=Guest("Template Test Guest", "guest@example.com", "+15550123456"),
        quote_breakdown={
            "line_items": display_items,
            "display_line_items": display_items,
            "subtotal_net": 10000.00,
            "total_vat": 1250.00,
            "property_fee": 500.00,
            "property_fee_rate": 0.05,
            "currency": "USD",
            "vat_by_class": {"accommodation": 500.00, "standard": 750.00},
            "package_title": "Template Test Package",
            "included_amenities": "Breakfast and tours",
            "check_in": "2026-12-01",
            "check_out": "2026-12-08",
            "total_guests": 4,
            "accommodation_share": 0.5,
            "promo_code": "",
            "promo_discount_rate": 0,
            "promo_discount_amount": 0,
            "payments": payments,
            "booking_number": "WC-TEST-001",
        },
    )


def all_docx_text(path):
    doc = Document(path)
    values = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def test_context_formats_invoice_and_balance():
    context = build_invoice_context(make_invoice(room_count=2, payment_count=2))
    assert context["invoice_number"] == "WBE-TEST-001"
    assert context["total_due"] == "$11,750.00"
    assert context["remaining_balance"] == "$11,450.00"
    assert context["rooms"][0] == {"name": "Room 1", "quantity": "1", "nights": "7"}
    assert context["payments"][1] == {"date": "2026-08-02", "amount": "$200.00"}


def test_context_accepts_exact_maximums():
    context = build_invoice_context(make_invoice(MAX_ROOMS, MAX_PAYMENTS))
    assert len(context["rooms"]) == 5
    assert len(context["payments"]) == 4


def test_context_rejects_more_than_five_rooms():
    with pytest.raises(ValueError, match="maximum of 5 rooms"):
        build_invoice_context(make_invoice(MAX_ROOMS + 1, 1))


def test_context_rejects_room_quantities_totaling_more_than_five():
    invoice = make_invoice(2, 1)
    invoice.lines[0].room_quantity = 3
    invoice.lines[1].room_quantity = 3
    with pytest.raises(ValueError, match="maximum of 5 rooms"):
        build_invoice_context(invoice)


def test_context_rejects_more_than_four_payments():
    with pytest.raises(ValueError, match="maximum of 4 payments"):
        build_invoice_context(make_invoice(1, MAX_PAYMENTS + 1))


def test_render_populates_uploaded_template_without_tokens(tmp_path):
    output = tmp_path / "rendered.docx"
    template = Path(__file__).resolve().parents[1] / "invoice_template.docx"
    render_invoice_docx(make_invoice(MAX_ROOMS, MAX_PAYMENTS), output, template)
    text = all_docx_text(output)
    assert output.exists()
    assert "WBE-TEST-001" in text
    assert all(f"Room {index}" in text for index in range(1, 6))
    assert all(f"2026-08-{index:02d}" in text for index in range(1, 5))
    assert "{{" not in text
    assert "{%" not in text


def test_render_word_pdf_creates_real_pdf(tmp_path):
    if not find_libreoffice():
        pytest.skip("LibreOffice is not installed")
    output = tmp_path / "WBE-TEST-001.pdf"
    result = render_invoice_word_pdf(make_invoice(MAX_ROOMS, MAX_PAYMENTS), output)
    assert result == str(output)
    assert output.exists()
    assert output.stat().st_size > 1000
    assert output.read_bytes().startswith(b"%PDF")
