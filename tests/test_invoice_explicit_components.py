"""Disconnected renderer projection; fixtures are arithmetic, not authority."""
import copy
import pytest
from datetime import date
from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from booking_engine.invoice import Guest, Invoice
from booking_engine.invoice_word import build_invoice_context


def fixture():
    # Two ORIGINAL q=1 groups: .13 each, never merged to .24.
    amounts = dict(grossCents=20, discountCents=0, roomTotalCents=20,
                   propertyFeeCents=2, accommodationVatCents=2,
                   packageVatCents=2, totalVatCents=4, grandTotalCents=26)
    lines = [dict(label=f"Original group {i}", quantity=2, room_quantity=1,
                  unit_price=.05, net=.10, vat_rate=0, vat=.02, gross=.12)
             for i in range(2)]
    quote = dict(line_items=lines, subtotal_net=.20, total_vat=.04,
                 property_fee=.02, total=.26, promo_discount_amount=0,
                 vat_by_class=dict(accommodation=.02, standard=.02))
    return quote, amounts


def invoice(quote, amounts):
    return Invoice.from_quote("OFFLINE-TEST", date(2026, 9, 6),
                              Guest("Local Fixture", "fixture@example.invalid", "+15555550100"),
                              quote, component_cents=amounts)


def rendered_pdf_text(inv, route, tmp_path, monkeypatch):
    from booking_engine.invoice_pdf import render_invoice_pdf
    from booking_engine import invoice_renderer
    import pymupdf
    output = tmp_path / f"{route}.pdf"
    if route == "reportlab":
        render_invoice_pdf(inv, str(output))
    else:
        monkeypatch.setenv("WBE_INVOICE_RENDERER", "word")
        monkeypatch.setenv("WBE_INVOICE_REPORTLAB_FALLBACK", "1" if route == "fallback" else "0")
        if route == "fallback":
            def fail_word(*args):
                raise RuntimeError("offline forced conversion failure")
            monkeypatch.setattr(invoice_renderer, "render_invoice_word_pdf", fail_word)
        used = invoice_renderer.render_invoice_pdf_for_service(inv, output)
        assert used == ("reportlab-fallback" if route == "fallback" else "word")
    with pymupdf.open(output) as doc:
        return "\n".join(page.get_text() for page in doc)


def assert_pdf_amount(text, label, amount):
    import re
    # Word may wrap a long amount inside its unchanged narrow template cell.
    exact_amount = r"\s*".join(re.escape(char) for char in amount)
    assert re.search(re.escape(label) + r"\s+" + exact_amount + r"(?=\s|$)", text), text


@pytest.mark.parametrize("route", ["reportlab", "fallback", "word"])
def test_accepted_decimal_total_renders_exact_pdf(route, tmp_path, monkeypatch):
    from decimal import Decimal
    quote, amounts = fixture()
    quote["total"] = Decimal("0.26")
    inv = invoice(quote, amounts)
    assert isinstance(inv.total, Decimal)  # Admission must not narrow to floats.
    assert build_invoice_context(inv)["total_due"] == "$0.26"
    text = rendered_pdf_text(inv, route, tmp_path, monkeypatch)
    assert_pdf_amount(text, "TOTAL DUE", "$0.26")
    assert_pdf_amount(text, "Remaining Balance:", "$0.26")


@pytest.mark.parametrize("route", ["reportlab", "fallback"])
@pytest.mark.parametrize("decimal_subtotal", [False, True])
def test_boundary_gross_uses_authoritative_cents(route, decimal_subtotal, tmp_path, monkeypatch):
    from decimal import Decimal
    quote, amounts = fixture()
    amounts.update(grossCents=4000000000000006, discountCents=5,
                   roomTotalCents=4000000000000001, propertyFeeCents=0,
                   accommodationVatCents=0, packageVatCents=0,
                   totalVatCents=0, grandTotalCents=4000000000000001)
    net = Decimal("40000000000000.01") if decimal_subtotal else 40000000000000.01
    quote.update(subtotal_net=net, promo_discount_amount=.05, total=net,
                 property_fee=0, total_vat=0,
                 vat_by_class=dict(accommodation=0, standard=0))
    quote["line_items"] = [dict(label="Large exact component", quantity=1,
        room_quantity=1, net=net, vat=0, gross=net, vat_rate=0)]
    inv = invoice(quote, amounts)
    assert build_invoice_context(inv)["subtotal_net"] == "$40,000,000,000,000.01"
    text = rendered_pdf_text(inv, route, tmp_path, monkeypatch)
    assert_pdf_amount(text, "Adventure Package Fees", "$40,000,000,000,000.06")
    assert_pdf_amount(text, "Promo discount", "-$0.05")
    assert_pdf_amount(text, "Subtotal after discount", "$40,000,000,000,000.01")
    assert_pdf_amount(text, "TOTAL DUE", "$40,000,000,000,000.01")
    assert_pdf_amount(text, "Remaining Balance:", "$40,000,000,000,000.01")


@pytest.mark.parametrize("route", ["reportlab", "fallback", "word"])
def test_boundary_payment_balance_is_exact(route, tmp_path, monkeypatch):
    from decimal import Decimal
    quote, amounts = fixture()
    amounts.update(grossCents=4000000000000001, discountCents=0,
                   roomTotalCents=4000000000000001, propertyFeeCents=0,
                   accommodationVatCents=0, packageVatCents=0,
                   totalVatCents=0, grandTotalCents=4000000000000001)
    # A float total also witnesses the old subtraction's wrong cent without a TypeError.
    quote.update(subtotal_net=40000000000000.01, total=40000000000000.01,
                 property_fee=0, total_vat=0,
                 vat_by_class=dict(accommodation=0, standard=0),
                 payments=[dict(datePaid="2026-09-01", paymentAmount=Decimal("0.02"))])
    quote["line_items"] = [dict(label="Large paid component", quantity=1,
        room_quantity=1, net=40000000000000.01, vat=0, gross=40000000000000.01, vat_rate=0)]
    inv = invoice(quote, amounts)
    text = rendered_pdf_text(inv, route, tmp_path, monkeypatch)
    assert_pdf_amount(text, "TOTAL DUE", "$40,000,000,000,000.01")
    assert_pdf_amount(text, "2026-09-01", "$0.02")
    assert_pdf_amount(text, "Remaining Balance:", "$39,999,999,999,999.99")


@pytest.mark.parametrize("route", ["reportlab", "fallback", "word"])
@pytest.mark.parametrize("complement", [False, True])
def test_mixed_decimal_money_labels_match_each_component(route, complement, tmp_path, monkeypatch):
    from decimal import Decimal
    quote, amounts = fractional_fixture()
    # Complementary mixes cover each admitted financial field as Decimal,
    # without changing the already accepted economics or payment conventions.
    fields = ("total_vat", "promo_discount_amount") if complement else ("subtotal_net", "property_fee", "total")
    for field in fields:
        quote[field] = Decimal(str(quote[field]))
    vat_class = "accommodation" if complement else "standard"
    quote["vat_by_class"][vat_class] = Decimal(str(quote["vat_by_class"][vat_class]))
    for field in (("vat",) if complement else ("net", "gross")):
        quote["line_items"][0][field] = Decimal(str(quote["line_items"][0][field]))
    quote["payments"] = [dict(datePaid="2026-09-01", paymentAmount=Decimal("0.01")),
                         dict(datePaid="2026-09-02", paymentAmount=.02),
                         dict(date="2026-09-03", amount="100.00")]
    inv = invoice(quote, amounts)
    context = build_invoice_context(inv)
    assert {key: context[key] for key in ("subtotal_net", "total_vat", "property_fee", "total_due",
            "accommodation_vat_formula", "services_vat_formula", "remaining_balance")} == {
        "subtotal_net": "$249.23", "total_vat": "$31.15", "property_fee": "$12.46",
        "total_due": "$292.84", "accommodation_vat_formula": "$12.46",
        "services_vat_formula": "$18.69", "remaining_balance": "$192.81"}
    text = rendered_pdf_text(inv, route, tmp_path, monkeypatch)
    assert_pdf_amount(text, "Adventure Package Fees", "$249.23" if route == "word" else "$276.92")
    if route != "word":  # Unchanged Word template has net only, no discount row.
        assert_pdf_amount(text, "Promo discount", "-$27.69")
        assert_pdf_amount(text, "Subtotal after discount", "$249.23")
    assert_pdf_amount(text, "Total VAT", "$31.15")
    assert_pdf_amount(text, "Property Fee" if route == "word" else "Property fee", "$12.46")
    assert_pdf_amount(text, "TOTAL DUE", "$292.84")
    assert_pdf_amount(text, "Subtotal:", "$249.23")
    assert_pdf_amount(text, "Accommodation VAT:", "$12.46")
    assert_pdf_amount(text, "Services VAT:", "$18.69")
    assert_pdf_amount(text, "Total VAT:", "$31.15")
    for day, amount in [("01", "$0.01"), ("02", "$0.02"), ("03", "$100.00")]:
        assert_pdf_amount(text, "2026-09-" + day, amount)
    assert_pdf_amount(text, "Remaining Balance:", "$192.81")


def test_original_group_components_reach_word_without_reconstruction():
    quote, amounts = fixture()
    inv = invoice(quote, amounts)
    context = build_invoice_context(inv)
    assert context["total_due"] == "$0.26"
    assert context["accommodation_vat_formula"] == "$0.02"
    assert context["services_vat_formula"] == "$0.02"
    assert [room["name"] for room in context["rooms"]] == ["Original group 0", "Original group 1"]


@pytest.mark.parametrize("key,value", [
    ("grandTotalCents", 24), ("grossCents", 21), ("totalVatCents", 3),
    ("propertyFeeCents", True), ("propertyFeeCents", -1),
    ("propertyFeeCents", 2.0), ("propertyFeeCents", "2"),
    ("propertyFeeCents", float("nan")), ("propertyFeeCents", 2**53),
    ("verified", True),
])
def test_explicit_projection_rejects_invalid_or_unreconciled_cents(key, value):
    quote, amounts = fixture()
    amounts[key] = value
    with pytest.raises(ValueError, match="components"):
        invoice(quote, amounts)


@pytest.mark.parametrize("field", ["subtotal_net", "total_vat", "property_fee", "total", "promo_discount_amount"])
def test_explicit_projection_rejects_contradictory_quote(field):
    quote, amounts = fixture()
    quote[field] = .99
    with pytest.raises(ValueError, match="components"):
        invoice(quote, amounts)


def test_explicit_projection_rejects_contradictory_lines():
    quote, amounts = fixture()
    quote["line_items"][0]["net"] = .11
    with pytest.raises(ValueError, match="components"):
        invoice(quote, amounts)


def test_explicit_projection_rechecks_mutable_invoice():
    quote, amounts = fixture()
    inv = invoice(quote, amounts)
    inv.total = .24
    with pytest.raises(ValueError, match="components"):
        build_invoice_context(inv)


@pytest.mark.parametrize("renderer", ["reportlab", "fallback"])
def test_real_reportlab_routes_render_explicit_components(renderer, monkeypatch, tmp_path):
    from booking_engine import invoice_renderer
    import fitz
    quote, amounts = fixture()
    inv = invoice(quote, amounts)
    monkeypatch.setenv("WBE_INVOICE_RENDERER", "reportlab" if renderer == "reportlab" else "word")
    monkeypatch.setenv("WBE_INVOICE_REPORTLAB_FALLBACK", "1")
    if renderer == "fallback":
        def fail_word(*args):
            raise RuntimeError("offline conversion failure")
        monkeypatch.setattr(invoice_renderer, "render_invoice_word_pdf", fail_word)
    output = tmp_path / "explicit.pdf"
    used = invoice_renderer.render_invoice_pdf_for_service(inv, output)
    assert used == ("reportlab" if renderer == "reportlab" else "reportlab-fallback")
    with fitz.open(output) as document:
        text = "\n".join(page.get_text() for page in document)
    assert "$0.26" in text
    assert "10%" not in text and "15%" not in text
    assert "Accommodation VAT:" in text and "Services VAT:" in text
    assert text.count("$0.02") >= 3
    assert "Original group 0" in text and "Original group 1" in text


@pytest.mark.parametrize("field", ["total", "property_fee", "promo_discount_amount", "vat_by_class"])
def test_explicit_mode_requires_supplied_components_without_defaults(field):
    quote, amounts = fixture()
    if field == "property_fee":
        amounts["propertyFeeCents"] = 0
        amounts["grandTotalCents"] = 24
        quote["total"] = .24
    del quote[field]
    with pytest.raises(ValueError, match="components"):
        invoice(quote, amounts)


def test_explicit_mode_does_not_even_evaluate_legacy_allocations():
    quote, amounts = fixture()
    inv = invoice(quote, amounts)
    inv.accommodation_allocation = object()
    inv.services_allocation = object()
    assert build_invoice_context(inv)["accommodation_vat_formula"] == "$0.02"
    from booking_engine.invoice_pdf import _dominica_vat_summary_elems
    from reportlab.lib.styles import getSampleStyleSheet
    styles = getSampleStyleSheet()
    _dominica_vat_summary_elems(inv, styles["Normal"], styles["Normal"])


def fractional_fixture():
    quote, amounts = fixture()
    amounts.update(grossCents=27692, discountCents=2769, roomTotalCents=24923,
                   propertyFeeCents=1246, accommodationVatCents=1246,
                   packageVatCents=1869, totalVatCents=3115, grandTotalCents=29284)
    quote.update(subtotal_net=249.23, total_vat=31.15, property_fee=12.46,
                 total=292.84, promo_discount_amount=27.69, promo_code="FIXTURE",
                 vat_by_class=dict(accommodation=12.46, standard=18.69))
    quote["line_items"] = [dict(label="Original Penthouse group", quantity=3,
        room_quantity=1, unit_price=249.23, net=249.23, vat_rate=0,
        vat=31.15, gross=280.38)]
    return quote, amounts


def test_explicit_reportlab_omits_unsupplied_percentage_claims(tmp_path):
    import pymupdf
    from booking_engine.invoice_pdf import render_invoice_pdf
    quote, amounts = fractional_fixture()
    output = tmp_path / "fractional.pdf"
    render_invoice_pdf(invoice(quote, amounts), str(output))
    with pymupdf.open(output) as doc:
        text = "\n".join(p.get_text() for p in doc)
    assert "$292.84" in text and "$12.46" in text and "$18.69" in text
    assert "(0%)" not in text and "0% off" not in text
    assert "$276.92" in text and "$27.69" in text


def test_explicit_vat_mapping_cannot_be_reconstructed_when_missing():
    quote, amounts = fixture()
    amounts.update(accommodationVatCents=0, packageVatCents=4)
    quote["line_items"][0].update(tax_class="accommodation", vat=0, gross=.10)
    quote["line_items"][1].update(vat=.04, gross=.14)
    quote["vat_by_class"] = None
    with pytest.raises(ValueError, match="components"):
        invoice(quote, amounts)


def test_explicit_construction_ignores_unused_legacy_allocations():
    quote, amounts = fixture()
    quote["accommodationShare"] = "not a current rate"
    assert build_invoice_context(invoice(quote, amounts))["total_due"] == "$0.26"


@pytest.mark.parametrize("field", ["subtotal_net", "total_vat", "property_fee", "total", "promo_discount_amount"])
@pytest.mark.parametrize("value", [None, "0.20", [], {}])
def test_malformed_quote_amount_types_raise_value_error(field, value):
    quote, amounts = fixture()
    quote[field] = value
    with pytest.raises(ValueError, match="components"):
        invoice(quote, amounts)


def test_snapshot_detached_immutable_and_keyword_only():
    import inspect
    quote, amounts = fixture()
    inv = invoice(quote, amounts)
    amounts["grandTotalCents"] = 24
    assert inv._component_cents["grandTotalCents"] == 26
    with pytest.raises(TypeError):
        inv._component_cents["grandTotalCents"] = 24
    assert inspect.signature(Invoice.from_quote).parameters["component_cents"].kind is inspect.Parameter.KEYWORD_ONLY
    assert "_component_cents" not in inspect.signature(Invoice).parameters
    assert build_invoice_context(inv)["total_due"] == "$0.26"


@pytest.mark.parametrize("flag", ["component_cents", "_component_cents", "verified", "accepted", "explicit_components"])
def test_public_quote_flags_do_not_activate_projection(flag):
    quote, amounts = fixture()
    quote[flag] = amounts if "component" in flag else True
    inv = Invoice.from_quote("LEGACY", date(2026, 9, 6),
                             Guest("Local", "local@example.invalid", "+15555550100"), quote)
    assert inv.explicit_vat_amounts() is None
    assert "10%" in build_invoice_context(inv)["accommodation_vat_formula"]


@pytest.mark.parametrize("damage", ["total", "subtotal_net", "property_fee", "promo_discount_amount", "total_vat", "vat_by_class", "line-net", "line-vat", "line-gross", "empty-lines", "currency"])
@pytest.mark.parametrize("route", ["docx", "reportlab", "fallback"])
def test_actual_renderers_reject_postconstruction_mutation(damage, route, tmp_path, monkeypatch):
    from booking_engine.invoice_word import render_invoice_docx
    from booking_engine.invoice_pdf import render_invoice_pdf
    from booking_engine import invoice_renderer
    quote, amounts = fixture()
    inv = invoice(quote, amounts)
    if damage.startswith("line-"):
        setattr(inv.lines[0], damage[5:], .99)
    elif damage == "empty-lines":
        inv.lines.clear()
    elif damage == "vat_by_class":
        inv.vat_by_class["standard"] = .99
    elif damage == "currency":
        inv.currency = "EUR"
    else:
        setattr(inv, damage, .99)
    output = tmp_path / ("bad.docx" if route == "docx" else "bad.pdf")
    monkeypatch.setenv("WBE_INVOICE_RENDERER", "word")
    monkeypatch.setenv("WBE_INVOICE_REPORTLAB_FALLBACK", "1")
    if route == "fallback":
        def fail_word(*args):
            raise RuntimeError("offline conversion failure")
        monkeypatch.setattr(invoice_renderer, "render_invoice_word_pdf", fail_word)
    render = {"docx": render_invoice_docx, "reportlab": render_invoice_pdf,
              "fallback": invoice_renderer.render_invoice_pdf_for_service}[route]
    with pytest.raises(ValueError, match="components"):
        render(inv, str(output))
    assert not output.exists()


@pytest.mark.parametrize("kind", ["grouped", "fractional", "changed-rates", "full-discount"])
@pytest.mark.parametrize("route", ["docx", "word", "reportlab", "fallback"])
def test_actual_artifacts_preserve_supplied_economics(kind, route, tmp_path, monkeypatch):
    from booking_engine.invoice_word import render_invoice_docx
    from booking_engine import invoice_renderer
    from docx import Document
    import pymupdf
    quote, amounts = fixture() if kind == "grouped" else fractional_fixture()
    if kind == "changed-rates":
        quote.update(accommodationShare=.91, property_fee_rate=.44, promo_discount_rate=.73)
        # A different accepted VAT split, not current percentages.
        amounts.update(accommodationVatCents=1000, packageVatCents=2115)
        quote["vat_by_class"] = dict(accommodation=10, standard=21.15)
    if kind == "full-discount":
        amounts.update(discountCents=27692, roomTotalCents=0, propertyFeeCents=0,
                       accommodationVatCents=0, packageVatCents=0, totalVatCents=0, grandTotalCents=0)
        quote.update(subtotal_net=0, property_fee=0, total_vat=0, total=0,
                     promo_discount_amount=276.92, vat_by_class=dict(accommodation=0, standard=0))
        quote["line_items"][0].update(net=0, vat=0, gross=0)
    inv = invoice(quote, amounts)
    output = tmp_path / ("actual.docx" if route == "docx" else "actual.pdf")
    if route == "docx":
        render_invoice_docx(inv, output)
        doc = Document(output)
        text = "\n".join([p.text for p in doc.paragraphs] +
                         [cell.text for table in doc.tables for row in table.rows for cell in row.cells])
    else:
        monkeypatch.setenv("WBE_INVOICE_RENDERER", "reportlab" if route == "reportlab" else "word")
        monkeypatch.setenv("WBE_INVOICE_REPORTLAB_FALLBACK", "1" if route == "fallback" else "0")
        if route == "fallback":
            def fail_word(*args):
                raise RuntimeError("offline conversion failure")
            monkeypatch.setattr(invoice_renderer, "render_invoice_word_pdf", fail_word)
        used = invoice_renderer.render_invoice_pdf_for_service(inv, output)
        assert used == ("reportlab-fallback" if route == "fallback" else route)
        with pymupdf.open(output) as doc:
            text = "\n".join(page.get_text() for page in doc)
    from decimal import Decimal
    for key in ("grandTotalCents", "propertyFeeCents", "accommodationVatCents", "packageVatCents"):
        assert f"${Decimal(amounts[key]) / 100:,.2f}" in text
    assert "10%" not in text and "15%" not in text and "44%" not in text
    positions = [text.index(line.label) for line in inv.lines]
    assert positions == sorted(positions)


@pytest.mark.parametrize("damage", ["missing-cent", "wrong-projection", "vat-class", "vat-total", "empty", "negative", "boolean", "fractional-cent", "infinity"])
def test_additional_malformed_projection_rejections(damage):
    quote, amounts = fixture()
    if damage == "missing-cent":
        del amounts["grossCents"]
    elif damage == "wrong-projection":
        amounts = []
    elif damage == "vat-class":
        quote["vat_by_class"]["other"] = 0
    elif damage == "vat-total":
        quote["vat_by_class"]["standard"] = .03
    elif damage == "empty":
        quote["line_items"] = []
    else:
        quote["subtotal_net"] = {"negative": -1, "boolean": True,
                                 "fractional-cent": .201, "infinity": float("inf")}[damage]
    with pytest.raises(ValueError, match="components"):
        invoice(quote, amounts)


def test_explicit_discount_is_shown_without_optional_promo_code(tmp_path):
    import pymupdf
    from booking_engine.invoice_pdf import render_invoice_pdf
    quote, amounts = fractional_fixture()
    del quote["promo_code"]
    output = tmp_path / "discount.pdf"
    render_invoice_pdf(invoice(quote, amounts), str(output))
    with pymupdf.open(output) as doc:
        text = "\n".join(p.get_text() for p in doc)
    assert "Promo discount" in text and "-$27.69" in text


@pytest.mark.parametrize("damage", ["missing-line-key", "no-lines", "non-list", "non-dict-line"])
def test_malformed_explicit_lines_raise_value_error(damage):
    quote, amounts = fixture()
    if damage == "missing-line-key":
        del quote["line_items"][0]["gross"]
    elif damage == "no-lines":
        del quote["line_items"]
    elif damage == "non-list":
        quote["line_items"] = None
    else:
        quote["line_items"] = [None]
    with pytest.raises(ValueError, match="components"):
        invoice(quote, amounts)



